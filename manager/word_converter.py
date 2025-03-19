from docx import Document
from docx.shared import Cm
import paragraphs

document = Document()

document.add_heading("ResumeBuilder", 0)

p = document.add_paragraph("Hello, this is a test documentument.")
p.add_run("This text is bold").bold = True
p.add_run("This text is bold").italic = True

document.add_paragraph("Experience 1", style="List Bullet")
document.add_paragraph("Experience 2", style="List Bullet")

margin = 1
sections = document.sections
for section in sections:
    section.top_margin = Cm(margin/2)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
    
document.save('test.docx')
